"""Shared docx parsing utilities."""

from docx import Document


def load_paragraphs(filepath: str) -> list[str]:
    """Load a docx and return a list of stripped paragraph texts."""
    doc = Document(filepath)
    return [p.text.strip() for p in doc.paragraphs]


def extract_images(filepath: str) -> dict[str, dict]:
    """Extract embedded images from a docx. Returns {ref: {blob, content_type}}."""
    doc = Document(filepath)
    images = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            images[rel.target_ref] = {
                "blob": rel.target_part.blob,
                "content_type": rel.target_part.content_type,
            }
    return images


def extract_description_by_range(
    paragraphs: list[str], start: int, end: int, max_length: int = 2000
) -> str:
    """Join paragraph texts in [start, end] range, skipping empties."""
    parts = [paragraphs[i] for i in range(start, min(end + 1, len(paragraphs))) if paragraphs[i]]
    desc = "\n\n".join(parts)
    if len(desc) > max_length:
        desc = desc[: max_length - 3] + "..."
    return desc


def extract_description_by_heading(
    paragraphs: list[str],
    heading_start: str,
    stop_headings: list[str],
    max_length: int = 2000,
) -> str:
    """Extract text starting from a heading until a stop condition is met."""
    parts = []
    for i, text in enumerate(paragraphs):
        if not text.startswith(heading_start):
            continue
        parts.append(text)
        for j in range(i + 1, len(paragraphs)):
            next_text = paragraphs[j]
            if not next_text or next_text.startswith("Section "):
                break
            if any(next_text.startswith(h) for h in stop_headings):
                break
            parts.append(next_text)
        break

    desc = "\n\n".join(parts)
    if len(desc) > max_length:
        desc = desc[: max_length - 3] + "..."
    return desc
